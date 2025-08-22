from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
from docx import Document
import re
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Number to words conversion function for Russian rubles
def number_to_words_ru(num):
    """Convert integer to Russian words for rubles"""
    if num == 0:
        return "ноль рублей"
    
    ones = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", 
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", 
            "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", 
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    
    thousands = ["", "тысяча", "тысячи", "тысяч"]
    millions = ["", "миллион", "миллиона", "миллионов"]
    
    def get_case_for_number(n, forms):
        """Get correct case for Russian numbers"""
        if n % 100 in [11, 12, 13, 14]:
            return forms[3]  # genitive plural for teens
        elif n % 10 == 1:
            return forms[1]  # nominative singular
        elif n % 10 in [2, 3, 4]:
            return forms[2]  # genitive singular
        else:
            return forms[3]  # genitive plural
    
    def convert_hundreds(n):
        """Convert numbers up to 999"""
        result = []
        
        h = n // 100
        if h > 0:
            result.append(hundreds[h])
        
        remainder = n % 100
        
        if remainder >= 10 and remainder < 20:
            result.append(teens[remainder - 10])
        else:
            t = remainder // 10
            if t > 0:
                result.append(tens[t])
            
            o = remainder % 10
            if o > 0:
                result.append(ones[o])
        
        return " ".join(result)
    
    if num >= 1000000:
        millions_part = num // 1000000
        millions_word = convert_hundreds(millions_part)
        millions_case = get_case_for_number(millions_part, millions)
        result = f"{millions_word} {millions_case}"
        
        remainder = num % 1000000
        if remainder > 0:
            result += " " + number_to_words_ru(remainder)
        else:
            # Add ruble case for millions without remainder
            ruble_forms = ["", "рубль", "рубля", "рублей"]
            ruble_case = get_case_for_number(num, ruble_forms)
            result += f" {ruble_case}"
        
        return result
    
    elif num >= 1000:
        thousands_part = num // 1000
        thousands_word = convert_hundreds(thousands_part)
        
        # Special case for feminine thousands
        if thousands_part % 10 == 1 and thousands_part % 100 != 11:
            thousands_word = thousands_word.replace("один", "одна")
        elif thousands_part % 10 == 2 and thousands_part % 100 != 12:
            thousands_word = thousands_word.replace("два", "две")
        
        thousands_case = get_case_for_number(thousands_part, thousands)
        result = f"{thousands_word} {thousands_case}"
        
        remainder = num % 1000
        if remainder > 0:
            result += " " + number_to_words_ru(remainder)
        else:
            # Add ruble case for thousands without remainder
            ruble_forms = ["", "рубль", "рубля", "рублей"]
            ruble_case = get_case_for_number(num, ruble_forms)
            result += f" {ruble_case}"
        
        return result
    
    else:
        # Numbers less than 1000
        word = convert_hundreds(num)
        
        # Add ruble case
        ruble_forms = ["", "рубль", "рубля", "рублей"]
        ruble_case = get_case_for_number(num, ruble_forms)
        
        return f"{word} {ruble_case}"

def generate_contract_number():
    """Generate contract number based on current date: КР + DD.MM.YY"""
    now = datetime.now(timezone.utc)
    return f"КР{now.strftime('%d.%m.%y')}"

def calculate_contract_end_date(duration_months: int):
    """Calculate contract end date based on duration"""
    now = datetime.now(timezone.utc)
    
    # Calculate end month
    end_month = now.month + duration_months
    end_year = now.year
    
    if end_month > 12:
        end_year = end_year + (end_month - 1) // 12
        end_month = ((end_month - 1) % 12) + 1
    
    # Get the last day of the end month
    if end_month == 12:
        next_month = 1
        next_year = end_year + 1
    else:
        next_month = end_month + 1
        next_year = end_year
    
    last_day_of_month = (datetime(next_year, next_month, 1) - timedelta(days=1)).day
    
    end_date = datetime(end_year, end_month, last_day_of_month)
    
    # Format for template
    months_ru = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
        7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    
    return str(end_date.day), months_ru[end_date.month], str(end_date.year)

# Contract template - the default contract
CONTRACT_TEMPLATE = """**Договор об оказании услуг № {contract_number}**

г. Казань «___» 2025 г.

Индивидуальный предприниматель Шамсутдинов Радис Раисович, именуемый в дальнейшем «Исполнитель» с одной стороны и {client_name}, именуемый в дальнейшем «Заказчик», с другой стороны, далее совместно именуемые «Стороны» заключили настоящий Договор о нижеследующем:

**1. ПРЕДМЕТ ДОГОВОРА**

1.1. «Исполнитель» принимает на себя обязательства оказать комплекс услуг в соответствии с заявками «Заказчика», а «Заказчик» обязуется принять услуги и оплатить их в размере и порядке, установленном настоящим договором.

1.2. В комплекс оказываемых услуг входят:

1.2.1. Создание рекламных кампаний в Яндекс.Директ.

Создание кампании контекстной рекламы включает в себя:

- Анализ поискового спроса по тематике деятельности, указанной Заказчиком;
- Подбор ключевых запросов, по которым будут размещаться рекламные объявления заказчика;
- Подготовка ключевых запросов к публикации в контекстной системе;
- Составление текстовых блоков объявлений, на основе информационных материалов, предоставленных заказчиком;
- Публикация ключевых запросов, текстовых блоков и веб-страниц в аккаунтах контекстных систем сети Интернет. Запуск рекламных кампаний;

1.2.2. В подарок «Исполнитель» осуществляет ведение рекламных кампаний в течение 3 (трёх) календарных недель после запуска. Ведение рекламной кампании по контекстной рекламе:

- Управление ценой клика кампаний;
- Мониторинг изменений позиций объявлений;
- Мониторинг эффективности текстовых блоков объявлений;
- Мониторинг статуса ключевых запросов;
- Мониторинг CTR кампаний;

1.3. Настройка и ведение рекламных кампаний осуществляются через аккаунты, созданные в системе Витамин. Данный сервис является партнером Яндекса и позволяет оптимизировать рекламные кампании. Например, одна из возможностей сервиса - автоматическое управление ставками, сэкономит вам бюджет и позволит получать клики по самой выгодной цене

**2. СРОК ДЕЙСТВИЯ ДОГОВОРА**

2.1. Настоящий Договор вступает в силу с даты его подписания Сторонами и действует до «{contract_end_date}» {contract_end_month} {contract_end_year} года.

2.2. Договор может быть расторгнут в одностороннем порядке по инициативе одной из Сторон при условии письменного уведомления другой Стороны, но не позднее чем за 7 (семь) дней до предполагаемой даты расторжения Договора.

2.3. Досрочное расторжение Договора возможно по взаимному согласию Сторон, выраженному в письменной форме.

2.4. Если иное не предусмотрено в соглашении сторон о досрочном расторжении договора, прекращение действия Договора не освобождает Стороны от необходимости исполнения всех своих обязательств, предусмотренных Договором, которые не были исполнены на момент прекращения его действия, а также не освобождает Стороны от ответственности за неисполнение (ненадлежащее исполнение) обязательств.

**3. ПРАВА И ОБЯЗАННОСТИ СТОРОН**

**3.1. «Исполнитель» обязан:**

3.1.1. Приступить к оказанию Услуг в течение трех дней с момента поступления оплаты за них.

3.1.2. Консультировать Заказчика по всем вопросам, касающихся предмета данного Договора.

3.1.3. Незамедлительно уведомлять «Заказчика» обо всех обстоятельствах, которые могут повлечь задержку в оказании Услуг.

3.1.4. Сохранять конфиденциальность условий настоящего Договора, а также информации, полученной от «Заказчика» в связи с исполнением настоящего Договора.

3.1.5. Предоставить доступы к статистике рекламных кампаний «Заказчика»».

3.1.6. До конца отчетного месяца предоставлять акт выполненных работ.

3.1.7. Направлять отчеты по запросу уполномоченного представителя в конце месяца оказания услуг.

3.1.8. Обязуется не допускать искажения предоставляемой «Заказчиком» для размещения (распространения) информации.

**3.2. «Исполнитель» вправе:**

3.2.1. Требовать от «Заказчика» предоставления необходимой информации для надлежащего оказания Услуг.

**3.3. «Заказчик» обязан:**

3.3.1. Предоставлять «Исполнителю» информацию, необходимую для оказания Услуг по настоящему Договору.

3.3.2. Оплатить Услуги в сроки и в порядке, установленные настоящим Договором.

3.3.3. Подписать Акт выполненных работ, предоставленный «Исполнителем», в течение 3 (трех) рабочих дней, либо предоставить претензии по Услугам с указанием перечня необходимых доработок и сроков их исполнения. «Исполнитель» обязуется устранить замечания своими силами и за свой счет. После устранения мотивированных возражений «Исполнитель» повторно направляет Акт выполненных работ согласно процедуре, описанной в настоящем пункте Договора. В случае просрочки «Заказчика» в подписании акта и или предоставлении претензий услуги считаются оказанными надлежащим образом и принятыми «Заказчиком» в полном объеме.

**3.4. «Заказчик» вправе:**

3.4.1. Проверять ход и качество оказываемых «Исполнителем» услуг.

3.4.2. Выдвигать требования, необходимые для надлежащего оказания Услуг.

**4. ЦЕНА УСЛУГ И ПОРЯДОК РАСЧЕТОВ**

4.1 Стоимость услуг по созданию рекламных кампаний составляет {service_cost} ({service_cost_words}) рублей в месяц.

4.2 Полная оплата производится в день подписания настоящего договора на расчетный счет Исполнителя. В дальнейшем Заказчик оплачивает услуги ежемесячно после выставления счета Исполнителем в течение трех дней.

**5. ПОРЯДОК СДАЧИ-ПРИЕМКИ УСЛУГ.**

5.1. Не позднее 3 (Три) рабочих дней с момента оказания услуг ежемесячно Стороны подписывают Акт выполненных работ (далее – Акт). С момента подписания обеими Сторонами Акта услуги считаются оказанными и принятыми Сторонами без возражений и замечаний.

5.2. Заказчик самостоятельно отслеживает получение акта, если в течение 03 рабочих дней после окончания периода оказания услуг акт не поступил на адрес Заказчика, то он незамедлительно сообщает об этом Исполнителю.

5.3. В том случае, если в течение 3 (Три) календарных дней с момента окончания периода оказания услуг Исполнитель не получит от Заказчика обоснованной претензии, услуги соответствующего периода будут признаваться оказанными надлежащим образом и принятыми Заказчиком в полном объеме, а Акт выполненных работ (услуг) за соответствующий период приобретает юридическую силу за подписью Исполнителя. Кроме того, безусловным подтверждением надлежащего оказания Услуг за период оказания услуг является оплата выставленного Исполнителем счета за Услуги данного периода и/или предоплата за Услуги следующего периода оказания услуг.

**6 ФОРС-МАЖОР**

6.1. Стороны освобождаются от ответственности за частичное или полное неисполнение своих обязанностей по настоящему Договору, если Сторона, для которой сложилась невозможность исполнения своих обязанностей, докажет, что данное неисполнение или ненадлежащее исполнение явилось следствием действия обстоятельств непреодолимой силы. Такими обстоятельствами считаются стихийные бедствия, вооруженные конфликты, забастовки, издание органами государственной власти и управления нормативных актов, препятствующих исполнению настоящего Договора, а также другие события, возникшие после подписания настоящего Договора и находящиеся вне разумного предвидения и контроля Сторон.

6.2. Обстоятельством непреодолимой силы признается также издание органами власти и управления актов, делающих невозможным исполнение обязательств по настоящему Договору хотя бы одной из Сторон.

**7. КОНФИДЕНЦИАЛЬНОСТЬ**

7.1. Любая информация, данные или сведения, полученные Сторонами в целях исполнения настоящего Договора, рассматриваются как конфиденциальные и не могут быть раскрыты третьим лицам, за исключением случаев, предусмотренных действующим законодательством Российской Федерации.

7.2. Конфиденциальная информация не может быть раскрыта третьим лицам, опубликована или другим образом разглашена в течение срока действия настоящего Договора, в случае отсутствия письменного разрешения Сторон.

**8. ОТВЕТСТВЕННОСТЬ СТОРОН**

8.1. За неисполнение, ненадлежащее исполнение своих обязательств по настоящему Договору Стороны несут ответственность в соответствии с действующим законодательством РФ.

**9. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ**

9.1. Все споры или разногласия, возникающие между Сторонами по настоящему Договору, разрешаются путем переговоров.

9.2. В случае недостижения соглашения в ходе переговоров заинтересованная сторона направляет заказным письмом претензию в письменной форме, подписанную уполномоченным лицом. Срок рассмотрения претензии – 3 (три) календарных дня с даты получения претензии другой стороной.

9.3. Если Сторонам не удастся достичь согласия, то любой спор, разногласие или требование, возникающие из данного Договора или касающиеся его нарушения, подлежат разрешению в соответствии с законодательством РФ.

**10. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ**

10.1. Настоящий Договор составлен в двух экземплярах, по одному экземпляру для каждой из Сторон.

10.2. После подписания настоящего Договора все предыдущие переговоры и переписка, связанная с его заключением, теряют силу.

10.3. Настоящий Договор может быть изменен и/или дополнен только документом, составленным в письменной форме, подписанным Сторонами.

**11. ЮРИДИЧЕСКИЕ АДРЕСА И БАНКОВСКИЕ РЕКВИЗИТЫ СТОРОН**

| **«Исполнитель»:**<br>Индивидуальный предприниматель<br>Шамсутдинов Радис Раисович<br>Юридический адрес организации<br>423040, Россия, Республика Татарстан,<br>Нурлатский р-н, г. Нурлат,<br>ул. им Р.С. Хамадеева, д. 9, кв. 8<br>ИНН 163205154150<br>ОГРНИП 319169000185092<br>Р/с 40802810700001303517<br>Банк АО «ТБанк»<br>Юридический адрес банка<br>127287, г. Москва, ул. Хуторская 2-я,<br>д.38А, стр. 26<br>К/с 30101810145250000974<br>ИНН банка 7710140679<br>БИК 044525974<br>________________/Шамсутдинов Р.Р. | **«Заказчик»:**<br>{client_details}<br><br><br><br><br><br><br><br><br>________________/{client_signature} |

"""

# Models
class ContractData(BaseModel):
    """Input data for creating a contract directly"""
    name_or_organization: str  # Client name/organization
    other_details: Optional[str] = None  # Other client details
    service_cost: int  # Cost as integer (e.g., 30000)
    duration_months: int = 6  # Duration in months (1, 6, 12)

class ContractNew(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    contract_number: str  # Auto-generated: КР + DD.MM.YY
    client_name: str
    client_details: str  # Full client details for contract
    service_cost: int
    service_cost_words: str  # Auto-generated from service_cost
    contract_start_date: str  # Current date
    contract_end_date: str
    contract_end_month: str
    contract_end_year: Optional[str] = "2025"  # Add year field with default for backward compatibility
    contract_content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Keep old models for backward compatibility
class Client(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name_or_organization: str  # Первое поле: имя/название организации
    other_details: Optional[str] = None  # Второе поле: другие данные
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ClientCreate(BaseModel):
    name_or_organization: str  # Первое поле: имя/название организации
    other_details: Optional[str] = None  # Второе поле: другие данные

class Contract(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    client_name: str
    service_cost: str
    service_cost_words: str
    contract_end_date: str
    contract_end_month: str
    contract_content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ContractCreate(BaseModel):
    client_id: str
    service_cost: str
    service_cost_words: str
    contract_end_date: str
    contract_end_month: str

def add_formatted_paragraph(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a paragraph with consistent Times New Roman 11pt formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    para.alignment = alignment
    return para

def create_word_contract(contract_data):
    """Create a Word document with the contract content"""
    doc = Document()
    
    # Set document margins - smaller margins to fit content in 3 pages
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font for entire document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Header with Kazan (left) and contract signing date (right)
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add "Казань" on the left
    left_run = header_para.add_run("Казань")
    left_run.font.name = 'Times New Roman'
    left_run.font.size = Pt(11)
    
    # Add spaces to push the date to the right
    spacer_run = header_para.add_run("\t" * 10)  # Use tabs for alignment
    
    # Add current date on the right
    current_date = datetime.now(timezone.utc)
    date_str = f'«{current_date.day}» {current_date.strftime("%B")} {current_date.year} г.'
    # Replace English month names with Russian
    months_map = {
        'January': 'января', 'February': 'февраля', 'March': 'марта', 'April': 'апреля',
        'May': 'мая', 'June': 'июня', 'July': 'июля', 'August': 'августа',
        'September': 'сентября', 'October': 'октября', 'November': 'ноября', 'December': 'декабря'
    }
    for eng_month, ru_month in months_map.items():
        date_str = date_str.replace(eng_month, ru_month)
    
    right_run = header_para.add_run(date_str)
    right_run.font.name = 'Times New Roman'
    right_run.font.size = Pt(11)
    
    # Empty line after header
    doc.add_paragraph()
    
    # Title
    add_formatted_paragraph(doc, 
        f"Договор об оказании услуг № {contract_data['contract_number']}", 
        bold=True, 
        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Contract parties
    add_formatted_paragraph(doc, 
        f'Индивидуальный предприниматель Шамсутдинов Радис Раисович, именуемый в дальнейшем «Исполнитель» с одной стороны и {contract_data["client_name"]}, именуемый в дальнейшем «Заказчик», с другой стороны, далее совместно именуемые «Стороны» заключили настоящий Договор о нижеследующем:', 
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Section 1 - Subject
    add_formatted_paragraph(doc, 
        "1. ПРЕДМЕТ ДОГОВОРА", 
        bold=True, 
        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Content sections with proper formatting
    
    add_formatted_paragraph(doc, '1.1. «Исполнитель» принимает на себя обязательства оказать комплекс услуг в соответствии с заявками «Заказчика», а «Заказчик» обязуется принять услуги и оплатить их в размере и порядке, установленном настоящим договором.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '1.2. В комплекс оказываемых услуг входят:', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '1.2.1. Создание рекламных кампаний в Яндекс.Директ.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, 'Создание кампании контекстной рекламы включает в себя:', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Service list - compact formatting
    services = [
        '- Анализ поискового спроса по тематике деятельности, указанной Заказчиком; - Подбор ключевых запросов, по которым будут размещаться рекламные объявления заказчика;',
        '- Подготовка ключевых запросов к публикации в контекстной системе; - Составление текстовых блоков объявлений, на основе информационных материалов, предоставленных заказчиком;',
        '- Публикация ключевых запросов, текстовых блоков и веб-страниц в аккаунтах контекстных систем сети Интернет. Запуск рекламных кампаний;'
    ]
    
    for service in services:
        add_formatted_paragraph(doc, service, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '1.2.2. В подарок «Исполнитель» осуществляет ведение рекламных кампаний в течение 3 (трёх) календарных недель после запуска. Ведение рекламной кампании по контекстной рекламе:', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Management services - compact
    management_text = '- Управление ценой клика кампаний; - Мониторинг изменений позиций объявлений; - Мониторинг эффективности текстовых блоков объявлений; - Мониторинг статуса ключевых запросов; - Мониторинг CTR кампаний;'
    add_formatted_paragraph(doc, management_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '1.3. Настройка и ведение рекламных кампаний осуществляются через аккаунты, созданные в системе Витамин. Данный сервис является партнером Яндекса и позволяет оптимизировать рекламные кампании. Например, одна из возможностей сервиса - автоматическое управление ставками, сэкономит вам бюджет и позволит получать клики по самой выгодной цене', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Section 2 - Term
    add_formatted_paragraph(doc, "2. СРОК ДЕЙСТВИЯ ДОГОВОРА", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, f'2.1. Настоящий Договор вступает в силу с даты его подписания Сторонами и действует до «{contract_data["contract_end_date"]}» {contract_data["contract_end_month"]} {contract_data["contract_end_year"]} года.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '2.2. Договор может быть расторгнут в одностороннем порядке по инициативе одной из Сторон при условии письменного уведомления другой Стороны, но не позднее чем за 7 (семь) дней до предполагаемой даты расторжения Договора.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '2.3. Досрочное расторжение Договора возможно по взаимному согласию Сторон, выраженному в письменной форме.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '2.4. Если иное не предусмотрено в соглашении сторон о досрочном расторжении договора, прекращение действия Договора не освобождает Стороны от необходимости исполнения всех своих обязательств, предусмотренных Договором, которые не были исполнены на момент прекращения его действия, а также не освобождает Стороны от ответственности за неисполнение (ненадлежащее исполнение) обязательств.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Section 3 - Rights and Obligations  
    add_formatted_paragraph(doc, "3. ПРАВА И ОБЯЗАННОСТИ СТОРОН", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, '3.1. «Исполнитель» обязан:', bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Compact executor obligations
    executor_text = '3.1.1. Приступить к оказанию Услуг в течение трех дней с момента поступления оплаты за них. 3.1.2. Консультировать Заказчика по всем вопросам, касающихся предмета данного Договора. 3.1.3. Незамедлительно уведомлять «Заказчика» обо всех обстоятельствах, которые могут повлечь задержку в оказании Услуг.'
    add_formatted_paragraph(doc, executor_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    executor_text2 = '3.1.4. Сохранять конфиденциальность условий настоящего Договора, а также информации, полученной от «Заказчика» в связи с исполнением настоящего Договора. 3.1.5. Предоставить доступы к статистике рекламных кампаний «Заказчика»». 3.1.6. До конца отчетного месяца предоставлять акт выполненных работ.'
    add_formatted_paragraph(doc, executor_text2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '3.1.7. Направлять отчеты по запросу уполномоченного представителя в конце месяца оказания услуг. 3.1.8. Обязуется не допускать искажения предоставляемой «Заказчиком» для размещения (распространения) информации.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '3.2. «Исполнитель» вправе:', bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '3.2.1. Требовать от «Заказчика» предоставления необходимой информации для надлежащего оказания Услуг.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '3.3. «Заказчик» обязан:', bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    client_obligations_text = '3.3.1. Предоставлять «Исполнителю» информацию, необходимую для оказания Услуг по настоящему Договору. 3.3.2. Оплатить Услуги в сроки и в порядке, установленные настоящим Договором.'
    add_formatted_paragraph(doc, client_obligations_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '3.3.3. Подписать Акт выполненных работ, предоставленный «Исполнителем», в течение 3 (трех) рабочих дней, либо предоставить претензии по Услугам с указанием перечня необходимых доработок и сроков их исполнения. «Исполнитель» обязуется устранить замечания своими силами и за свой счет. После устранения мотивированных возражений «Исполнитель» повторно направляет Акт выполненных работ согласно процедуре, описанной в настоящем пункте Договора. В случае просрочки «Заказчика» в подписании акта и или предоставлении претензий услуги считаются оказанными надлежащим образом и принятыми «Заказчиком» в полном объеме.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, '3.4. «Заказчик» вправе:', bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '3.4.1. Проверять ход и качество оказываемых «Исполнителем» услуг. 3.4.2. Выдвигать требования, необходимые для надлежащего оказания Услуг.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Section 4 - Price and Payment
    add_formatted_paragraph(doc, "4. ЦЕНА УСЛУГ И ПОРЯДОК РАСЧЕТОВ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, f'4.1 Стоимость услуг по созданию рекламных кампаний составляет {contract_data["service_cost"]} ({contract_data["service_cost_words"]}) рублей в месяц.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_formatted_paragraph(doc, '4.2 Полная оплата производится в день подписания настоящего договора на расчетный счет Исполнителя. В дальнейшем Заказчик оплачивает услуги ежемесячно после выставления счета Исполнителем в течение трех дней.', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Remaining sections - compact
    add_formatted_paragraph(doc, "5. ПОРЯДОК СДАЧИ-ПРИЕМКИ УСЛУГ.", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "5.1. Не позднее 3 (Три) рабочих дней с момента оказания услуг ежемесячно Стороны подписывают Акт выполненных работ (далее – Акт). С момента подписания обеими Сторонами Акта услуги считаются оказанными и принятыми Сторонами без возражений и замечаний.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, "6 ФОРС-МАЖОР", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "6.1. Стороны освобождаются от ответственности за частичное или полное неисполнение своих обязанностей по настоящему Договору, если Сторона, для которой сложилась невозможность исполнения своих обязанностей, докажет, что данное неисполнение или ненадлежащее исполнение явилось следствием действия обстоятельств непреодолимой силы.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, "7. КОНФИДЕНЦИАЛЬНОСТЬ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "7.1. Любая информация, данные или сведения, полученные Сторонами в целях исполнения настоящего Договора, рассматриваются как конфиденциальные и не могут быть раскрыты третьим лицам, за исключением случаев, предусмотренных действующим законодательством Российской Федерации.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, "8. ОТВЕТСТВЕННОСТЬ СТОРОН", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "8.1. За неисполнение, ненадлежащее исполнение своих обязательств по настоящему Договору Стороны несут ответственность в соответствии с действующим законодательством РФ.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, "9. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "9.1. Все споры или разногласия, возникающие между Сторонами по настоящему Договору, разрешаются путем переговоров.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    add_formatted_paragraph(doc, "10. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "10.1. Настоящий Договор составлен в двух экземплярах, по одному экземпляру для каждой из Сторон.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Page break before section 11
    doc.add_page_break()
    
    # Section 11 - Details and signatures table (on new page)
    add_formatted_paragraph(doc, "11. ЮРИДИЧЕСКИЕ АДРЕСА И БАНКОВСКИЕ РЕКВИЗИТЫ СТОРОН", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Create table for signatures with equal height
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for column in table.columns:
        column.width = Inches(3)
    
    # Executor cell
    executor_cell = table.cell(0, 0)
    executor_cell.paragraphs[0].clear()  # Clear existing paragraph
    
    exec_para = executor_cell.add_paragraph()
    exec_run = exec_para.add_run("«Исполнитель»:")
    exec_run.bold = True
    exec_run.font.name = 'Times New Roman'
    exec_run.font.size = Pt(11)
    
    executor_details = [
        "Индивидуальный предприниматель",
        "Шамсутдинов Радис Раисович",  
        "Юридический адрес организации",
        "423040, Россия, Республика Татарстан,",
        "Нурлатский р-н, г. Нурлат,",
        "ул. им Р.С. Хамадеева, д. 9, кв. 8",
        "ИНН 163205154150",
        "ОГРНИП 319169000185092",
        "Р/с 40802810700001303517",
        "Банк АО «ТБанк»",
        "Юридический адрес банка",
        "127287, г. Москва, ул. Хуторская 2-я,",
        "д.38А, стр. 26",
        "К/с 30101810145250000974",
        "ИНН банка 7710140679",
        "БИК 044525974"
    ]
    
    for detail in executor_details:
        exec_detail_para = executor_cell.add_paragraph()
        exec_detail_run = exec_detail_para.add_run(detail)
        exec_detail_run.font.name = 'Times New Roman'
        exec_detail_run.font.size = Pt(11)
        # Make paragraph spacing more compact
        exec_detail_para.space_after = Pt(0)
        exec_detail_para.space_before = Pt(0)
    
    # Add signature line for executor
    exec_sig_para = executor_cell.add_paragraph()
    exec_sig_para.add_run("")  # Empty space
    exec_sig_para = executor_cell.add_paragraph()
    exec_sig_run = exec_sig_para.add_run("________________/Шамсутдинов Р.Р.")
    exec_sig_run.font.name = 'Times New Roman'
    exec_sig_run.font.size = Pt(11)
    
    # Client cell  
    client_cell = table.cell(0, 1)
    client_cell.paragraphs[0].clear()  # Clear existing paragraph
    
    client_para = client_cell.add_paragraph()
    client_run = client_para.add_run("«Заказчик»:")
    client_run.bold = True
    client_run.font.name = 'Times New Roman'
    client_run.font.size = Pt(11)
    
    # Format client details from contract data
    client_details_text = contract_data.get('client_details', contract_data['client_name'])
    client_details_lines = client_details_text.split('\n')
    
    for detail in client_details_lines:
        client_detail_para = client_cell.add_paragraph()
        client_detail_run = client_detail_para.add_run(detail)
        client_detail_run.font.name = 'Times New Roman'
        client_detail_run.font.size = Pt(11)
    
    # Add empty lines to match executor height
    for _ in range(max(0, len(executor_details) - len(client_details_lines))):
        client_cell.add_paragraph("")
    
    # Add signature line for client (on same level as executor)
    client_sig_para = client_cell.add_paragraph()
    client_sig_para.add_run("")  # Empty space
    client_sig_para = client_cell.add_paragraph()
    client_sig_run = client_sig_para.add_run(f"________________/{contract_data['client_name']}")
    client_sig_run.font.name = 'Times New Roman' 
    client_sig_run.font.size = Pt(11)
    
    return doc

# Helper function to convert MongoDB documents
def prepare_for_mongo(data):
    if isinstance(data.get('created_at'), datetime):
        data['created_at'] = data['created_at'].isoformat()
    return data

def parse_from_mongo(item):
    if isinstance(item.get('created_at'), str):
        item['created_at'] = datetime.fromisoformat(item['created_at'])
    return item

# API Routes
@api_router.get("/")
async def root():
    return {"message": "Contract Management System API"}

# Client endpoints
@api_router.post("/clients", response_model=Client)
async def create_client(client: ClientCreate):
    client_dict = client.dict()
    client_obj = Client(**client_dict)
    client_data = prepare_for_mongo(client_obj.dict())
    await db.clients.insert_one(client_data)
    return client_obj

@api_router.get("/clients", response_model=List[Client])
async def get_clients():
    clients = await db.clients.find().to_list(1000)
    result_clients = []
    
    for client in clients:
        client_parsed = parse_from_mongo(client)
        
        # Handle backward compatibility - convert old format to new format
        if 'name' in client_parsed and 'name_or_organization' not in client_parsed:
            # Old format - convert to new format
            name_or_org = client_parsed.get('name', '')
            if client_parsed.get('organization'):
                name_or_org = client_parsed.get('organization')
            
            other_details_parts = []
            if client_parsed.get('name') and client_parsed.get('organization'):
                other_details_parts.append(client_parsed.get('name'))
            if client_parsed.get('address'):
                other_details_parts.append(client_parsed.get('address'))
            if client_parsed.get('inn'):
                other_details_parts.append(f"ИНН {client_parsed.get('inn')}")
            if client_parsed.get('phone'):
                other_details_parts.append(f"Тел.: {client_parsed.get('phone')}")
            if client_parsed.get('email'):
                other_details_parts.append(f"Email: {client_parsed.get('email')}")
            
            client_data = {
                'id': client_parsed.get('id'),
                'name_or_organization': name_or_org,
                'other_details': '\n'.join(other_details_parts) if other_details_parts else None,
                'created_at': client_parsed.get('created_at')
            }
        else:
            # New format - use as is
            client_data = client_parsed
            
        result_clients.append(Client(**client_data))
    
    return result_clients

@api_router.get("/clients/{client_id}", response_model=Client)
async def get_client(client_id: str):
    client = await db.clients.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return Client(**parse_from_mongo(client))

@api_router.put("/clients/{client_id}", response_model=Client)
async def update_client(client_id: str, client_update: ClientCreate):
    existing_client = await db.clients.find_one({"id": client_id})
    if not existing_client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    updated_data = client_update.dict()
    updated_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.clients.update_one(
        {"id": client_id}, 
        {"$set": prepare_for_mongo(updated_data)}
    )
    
    updated_client = await db.clients.find_one({"id": client_id})
    return Client(**parse_from_mongo(updated_client))

@api_router.delete("/clients/{client_id}")
async def delete_client(client_id: str):
    result = await db.clients.delete_one({"id": client_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    return {"message": "Client deleted successfully"}

# New direct contract creation endpoint  
@api_router.post("/contracts/direct", response_model=ContractNew)
async def create_contract_direct(contract_data: ContractData):
    """Create contract directly without separate client record"""
    
    # Generate contract number
    contract_number = generate_contract_number()
    
    # Convert service cost to words
    service_cost_words = number_to_words_ru(contract_data.service_cost)
    
    # Calculate contract end date
    end_date, end_month, end_year = calculate_contract_end_date(contract_data.duration_months)
    
    # Format client details for contract
    client_details = contract_data.name_or_organization
    if contract_data.other_details:
        client_details += f"\n{contract_data.other_details}"
    
    # Use first field as client name in contract text
    client_name_in_contract = contract_data.name_or_organization
    
    # Generate contract content
    contract_content = CONTRACT_TEMPLATE.format(
        contract_number=contract_number,
        client_name=client_name_in_contract,
        service_cost=contract_data.service_cost,
        service_cost_words=service_cost_words,
        contract_end_date=end_date,
        contract_end_month=end_month,
        contract_end_year=end_year,
        client_details=client_details,
        client_signature=client_name_in_contract
    )
    
    # Create contract object
    now = datetime.now(timezone.utc)
    contract_obj = ContractNew(
        contract_number=contract_number,
        client_name=client_name_in_contract,
        client_details=client_details,
        service_cost=contract_data.service_cost,
        service_cost_words=service_cost_words,
        contract_start_date=now.strftime('%d.%m.%Y'),
        contract_end_date=end_date,
        contract_end_month=end_month,
        contract_end_year=end_year,
        contract_content=contract_content
    )
    
    # Save to database
    contract_data_dict = prepare_for_mongo(contract_obj.dict())
    await db.contracts_new.insert_one(contract_data_dict)
    
    return contract_obj

# Contract endpoints
@api_router.get("/contracts/direct", response_model=List[ContractNew])
async def get_contracts_direct():
    """Get all direct contracts"""
    contracts = await db.contracts_new.find().to_list(1000)
    return [ContractNew(**parse_from_mongo(contract)) for contract in contracts]

@api_router.get("/contracts/direct/{contract_id}", response_model=ContractNew)
async def get_contract_direct(contract_id: str):
    """Get specific direct contract"""
    contract = await db.contracts_new.find_one({"id": contract_id})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    return ContractNew(**parse_from_mongo(contract))

@api_router.put("/contracts/direct/{contract_id}", response_model=ContractNew)
async def update_contract_direct(contract_id: str, contract_data: ContractData):
    """Update direct contract with new data"""
    existing_contract = await db.contracts_new.find_one({"id": contract_id})
    if not existing_contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    # Keep original contract number and created_at
    existing_parsed = ContractNew(**parse_from_mongo(existing_contract))
    
    # Convert service cost to words
    service_cost_words = number_to_words_ru(contract_data.service_cost)
    
    # Calculate contract end date
    end_date, end_month, end_year = calculate_contract_end_date(contract_data.duration_months)
    
    # Format client details for contract
    client_details = contract_data.name_or_organization
    if contract_data.other_details:
        client_details += f"\n{contract_data.other_details}"
    
    # Use first field as client name in contract text
    client_name_in_contract = contract_data.name_or_organization
    
    # Generate updated contract content
    contract_content = CONTRACT_TEMPLATE.format(
        contract_number=existing_parsed.contract_number,
        client_name=client_name_in_contract,
        service_cost=contract_data.service_cost,
        service_cost_words=service_cost_words,
        contract_end_date=end_date,
        contract_end_month=end_month,
        contract_end_year=end_year,
        client_details=client_details,
        client_signature=client_name_in_contract
    )
    
    # Update contract
    updated_data = {
        "client_name": client_name_in_contract,
        "client_details": client_details,
        "service_cost": contract_data.service_cost,
        "service_cost_words": service_cost_words,
        "contract_end_date": end_date,
        "contract_end_month": end_month,
        "contract_end_year": end_year,
        "contract_content": contract_content,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.contracts_new.update_one(
        {"id": contract_id}, 
        {"$set": prepare_for_mongo(updated_data)}
    )
    
    updated_contract = await db.contracts_new.find_one({"id": contract_id})
    return ContractNew(**parse_from_mongo(updated_contract))

@api_router.delete("/contracts/direct/{contract_id}")
async def delete_contract_direct(contract_id: str):
    """Delete direct contract"""
    result = await db.contracts_new.delete_one({"id": contract_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Contract not found")
    return {"message": "Contract deleted successfully"}

@api_router.get("/contracts/direct/{contract_id}/download")
async def download_contract_direct_word(contract_id: str):
    """Download direct contract as Word document"""
    # Get contract from database
    contract = await db.contracts_new.find_one({"id": contract_id})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    contract_obj = ContractNew(**parse_from_mongo(contract))
    
    # Prepare contract data for Word generation
    contract_data = {
        "contract_number": contract_obj.contract_number,
        "client_name": contract_obj.client_name,
        "service_cost": str(contract_obj.service_cost),
        "service_cost_words": contract_obj.service_cost_words,
        "contract_end_date": contract_obj.contract_end_date,
        "contract_end_month": contract_obj.contract_end_month,
        "contract_end_year": contract_obj.contract_end_year,
        "client_details": contract_obj.client_details
    }
    
    # Create Word document
    doc = create_word_contract(contract_data)
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Create filename with better transliteration
    safe_client_name = contract_obj.client_name.replace(' ', '_')
    
    # More comprehensive transliteration
    cyrillic_to_latin = {
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e', 'ж': 'zh', 'з': 'z',
        'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm', 'н': 'n', 'о': 'o', 'п': 'p',
        'р': 'r', 'с': 's', 'т': 't', 'у': 'u', 'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch',
        'ш': 'sh', 'щ': 'sch', 'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
        'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'E', 'Ж': 'Zh', 'З': 'Z',
        'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M', 'Н': 'N', 'О': 'O', 'П': 'P',
        'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U', 'Ф': 'F', 'Х': 'H', 'Ц': 'Ts', 'Ч': 'Ch',
        'Ш': 'Sh', 'Щ': 'Sch', 'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya'
    }
    
    for cyrillic, latin in cyrillic_to_latin.items():
        safe_client_name = safe_client_name.replace(cyrillic, latin)
    
    # Remove any remaining non-ASCII characters and special characters
    safe_client_name = ''.join(c for c in safe_client_name if c.isalnum() or c in ['_', '-'])
    
    # Ensure filename is not empty
    if not safe_client_name:
        safe_client_name = "Contract"
    
    # Create safe contract number for filename
    safe_contract_number = contract_obj.contract_number.replace('.', '_').replace(' ', '_')
    safe_contract_number = ''.join(c for c in safe_contract_number if c.isalnum() or c in ['_', '-'])
    
    filename = f"Dogovor_{safe_client_name}_{safe_contract_number}.docx"
    
    # Use filename* parameter for better Unicode support
    from urllib.parse import quote
    filename_encoded = quote(filename.encode('utf-8'))
    
    return StreamingResponse(
        io.BytesIO(doc_buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}"}
    )

# Contract endpoints
@api_router.post("/contracts", response_model=Contract)
async def create_contract(contract: ContractCreate):
    # Get client details
    client = await db.clients.find_one({"id": contract.client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    client_obj = Client(**parse_from_mongo(client))
    
    # Format client details for contract - use the new structure
    client_details = client_obj.name_or_organization
    if client_obj.other_details:
        client_details += f"\n{client_obj.other_details}"
    
    # For client name in contract text, use the first field
    client_name_in_contract = client_obj.name_or_organization
    
    # Generate contract content
    contract_content = CONTRACT_TEMPLATE.format(
        contract_number="[Номер договора]",  # Legacy contracts don't have auto-generated numbers
        client_name=client_name_in_contract,
        service_cost=contract.service_cost,
        service_cost_words=contract.service_cost_words,
        contract_end_date=contract.contract_end_date,
        contract_end_month=contract.contract_end_month,
        contract_end_year="2025",  # Default year for legacy contracts
        client_details=client_details,
        client_signature=client_name_in_contract
    )
    
    # Create contract object
    contract_dict = contract.dict()
    contract_dict["client_name"] = client_name_in_contract
    contract_dict["contract_content"] = contract_content
    contract_obj = Contract(**contract_dict)
    
    # Save to database
    contract_data = prepare_for_mongo(contract_obj.dict())
    await db.contracts.insert_one(contract_data)
    
    return contract_obj

@api_router.get("/contracts", response_model=List[Contract])
async def get_contracts():
    contracts = await db.contracts.find().to_list(1000)
    return [Contract(**parse_from_mongo(contract)) for contract in contracts]

@api_router.get("/contracts/{contract_id}", response_model=Contract)
async def get_contract(contract_id: str):
    contract = await db.contracts.find_one({"id": contract_id})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    return Contract(**parse_from_mongo(contract))

@api_router.delete("/contracts/{contract_id}")
async def delete_contract(contract_id: str):
    result = await db.contracts.delete_one({"id": contract_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Contract not found")
    return {"message": "Contract deleted successfully"}

@api_router.get("/contracts/{contract_id}/download")
async def download_contract_word(contract_id: str):
    # Get contract from database
    contract = await db.contracts.find_one({"id": contract_id})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    contract_obj = Contract(**parse_from_mongo(contract))
    
    # Get client details for formatting
    client = await db.clients.find_one({"id": contract_obj.client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    client_obj = Client(**parse_from_mongo(client))
    
    # Format client details for Word document - use new structure
    client_details = client_obj.name_or_organization
    if client_obj.other_details:
        client_details += f"\n{client_obj.other_details}"
    
    # Prepare contract data for Word generation
    contract_data = {
        "contract_number": "[Номер договора]",  # Legacy contracts don't have auto-generated numbers
        "client_name": contract_obj.client_name,
        "service_cost": contract_obj.service_cost,
        "service_cost_words": contract_obj.service_cost_words,
        "contract_end_date": contract_obj.contract_end_date,
        "contract_end_month": contract_obj.contract_end_month,
        "contract_end_year": "2025",  # Default year for legacy contracts
        "client_details": client_details
    }
    
    # Create Word document
    doc = create_word_contract(contract_data)
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Create filename
    # Use transliteration or ASCII-safe filename to avoid encoding issues
    safe_client_name = contract_obj.client_name.replace(' ', '_').replace('ё', 'e').replace('Ё', 'E')
    # Simple transliteration for common Cyrillic characters
    cyrillic_to_latin = {
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ж': 'zh', 'з': 'z',
        'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm', 'н': 'n', 'о': 'o', 'п': 'p',
        'р': 'r', 'с': 's', 'т': 't', 'у': 'u', 'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch',
        'ш': 'sh', 'щ': 'sch', 'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
        'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ж': 'Zh', 'З': 'Z',
        'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M', 'Н': 'N', 'О': 'O', 'П': 'P',
        'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U', 'Ф': 'F', 'Х': 'H', 'Ц': 'Ts', 'Ч': 'Ch',
        'Ш': 'Sh', 'Щ': 'Sch', 'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya'
    }
    
    for cyrillic, latin in cyrillic_to_latin.items():
        safe_client_name = safe_client_name.replace(cyrillic, latin)
    
    filename = f"Dogovor_{safe_client_name}_{contract_obj.id[:8]}.docx"
    
    return StreamingResponse(
        io.BytesIO(doc_buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()